#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Gau2Publish
Author: Raúl Losantos
---------------------------------
Convert Gaussian (.log) outputs into:
 - .geom : ordered atomic coordinates with metadata (charge, multiplicity, energies, frequencies)
 - .png  : 3D render via `xyzrender` (if available)
 - .cdxml: ChemDraw document per molecule (if `pycdxml` is available)
 - DOCX  : document with title + 3D image + geometry block
 - ALL.cdxml: mosaic with thumbnails and properties

Highlights:
 - `-v/--verbose`: keep intermediates (.geom, .png, .mol2, .cdxml) and increase logging
 - `--xyzrender-args`: forward arbitrary flags to `xyzrender` (e.g., `-I`)
 - If `-I` is present in `--xyzrender-args`, files matching `*.v000.xyz` are removed after rendering
 - `--no-xyzrender`: skip 3D rendering (no PNG generated)
 - `--rebuild`: build DOCX only from existing .geom and .png (same basenames); do not run xyzrender or generate CDXML; never delete PNGs
 - `--scf-index=base,high`: define indices used to pick energies from `scfenergies`
   * Distances from the end: D_base = base + 1; D_high = (base + 1) - high
   * Python indices: idx_base = -D_base; idx_high = -D_high
   * Default: idx_base = -1; idx_high = -1
By default (without `-v`) the intermediates `.mol2`, `.geom`, `.png`, and `.cdxml` are deleted after being embedded into the DOCX/slide.
"""
from __future__ import annotations

# Imports
import argparse
import glob
import os
import shlex
import subprocess
import sys
from pathlib import Path
from typing import Tuple

# Scientific tools
import numpy as np
import pandas as pd
import cclib
from periodictable import elements
from rdkit import Chem
from rdkit.Chem.rdchem import BondDir, BondStereo, ChiralType

# Word export
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# pycdxml (optional)
try:
    from pycdxml import cdxml_converter, cdxml_slide_generator
    HAVE_PYCDXML = True
except Exception:
    HAVE_PYCDXML = False

from . import __version__

# ──────────────────────────────────────────────
# Units and helpers
EV_TO_HARTREE = 27.21138505

def eV_to_hartree(eV_value: float) -> float:
    return float(eV_value) / EV_TO_HARTREE

# ──────────────────────────────────────────────
# --scf-index parsing

def parse_scf_index(index: str | None, scfenergies_len: int) -> Tuple[int, int]:
    """Return (idx_base, idx_high) as negative Python indices.
    If `index` is None -> (-1, -1).
    Format: "base,high" with integers >= 0.
    Validates ranges and available length.
    """
    if scfenergies_len is None or scfenergies_len <= 0:
        raise RuntimeError("The file does not contain 'scfenergies'.")
    if index is None:
        return -1, -1
    try:
        s_base, s_high = index.split(',')
        base = int(s_base.strip())
        high = int(s_high.strip())
    except Exception as exc:
        raise ValueError("Invalid --scf-index format: use --scf-index=base,high (two integers)") from exc
    if base < 0 or high < 0:
        raise ValueError("--scf-index requires integers >= 0")
    D_base = base + 1
    D_high = (base + 1) - high
    if D_high <= 0:
        raise ValueError("--scf-index: constraint 0 ≤ high ≤ base+1 (so that D_high ≥ 1)")
    if max(D_base, D_high) > scfenergies_len:
        raise ValueError(
            f"--scf-index requests D_base={D_base}, D_high={D_high}, but scfenergies only has {scfenergies_len} entries"
        )
    return -D_base, -D_high

# ──────────────────────────────────────────────
# Gaussian parsing via cclib

def read_gaussian_output(gaussian_output_file: str, *, scf_index: str | None = None):
    """Parse a Gaussian .log file and return:
    charge, mult, scf_energy, scf_energy_high,
    enthalpy, free_energy, free_energy_corr,
    frequencies, atomic_numbers, geometry
    """
    try:
        data = cclib.io.ccread(gaussian_output_file)
    except Exception as exc:
        raise RuntimeError(f"Error reading Gaussian: {exc}") from exc
    if data is None:
        raise RuntimeError("cclib returned no data for the provided file.")

    charge = int(getattr(data, "charge", 0))
    mult = int(getattr(data, "mult", 1))

    scfenergies = getattr(data, "scfenergies", None)
    if scfenergies is None or len(scfenergies) == 0:
        raise RuntimeError("The file does not contain 'scfenergies'.")

    idx_base, idx_high = parse_scf_index(scf_index, len(scfenergies))
    scf_energy = eV_to_hartree(scfenergies[idx_base])
    scf_energy_high = eV_to_hartree(scfenergies[idx_high])

    enthalpy = getattr(data, "enthalpy", None)
    free_energy = getattr(data, "freeenergy", None)
    free_energy_corr = (float(free_energy) - float(scf_energy)) if free_energy is not None else None

    frequencies = list(getattr(data, "vibfreqs", []))
    atomic_numbers = list(map(int, getattr(data, "atomnos", [])))
    geometry = np.array(getattr(data, "atomcoords", [[0, 0, 0]]))[-1]
    
    return (
        charge, mult, scf_energy, scf_energy_high,
        enthalpy, free_energy, free_energy_corr,
        frequencies, atomic_numbers, geometry,
    )

# ──────────────────────────────────────────────
# Chemical symbols

def atomic_numbers_to_symbols(atomic_numbers: list[int]) -> list[str]:
    from periodictable import elements
    return [elements[n].symbol for n in atomic_numbers]

# ──────────────────────────────────────────────
# .geom writing

def write_geom_file(
    charge, mult, scf_energy, scf_energy_high, free_energy_corr,
    frequencies, atomic_numbers, geometry, method: str, out_base: Path
) -> str:
    """Write coordinates + metadata to a .geom file (sorted by descending Z)."""
    order = np.argsort(atomic_numbers)[::-1]
    sorted_coords = [geometry[i] for i in order]
    sorted_znums = [atomic_numbers[i] for i in order]
    symbols = atomic_numbers_to_symbols(sorted_znums)
    # Avoid repeated High energy method if is the standardly used
    if scf_energy == scf_energy_high:
        header = [
            f"{charge} {mult}",
            f"SCF Energy: {scf_energy:.6f} h",
        ]
    else:
        header = [
            f"{charge} {mult}",
            f"SCF Energy: {scf_energy:.6f} h",
            f"SCF Energy {method}: {scf_energy_high:.6f} h",
        ]

    if free_energy_corr is not None:
        header.append(f"Gcorr: {free_energy_corr:.6f} h")
    if len(frequencies) > 0:
        header.append(
            f"Imaginary Frequencies: {frequencies[0]:.2f} cm^-1" if float(frequencies[0]) < 0 else "No imaginary frequencies"
        )
    header.append("")

    df = pd.DataFrame({
        " ": symbols,
        "x": [float(c[0]) for c in sorted_coords],
        "y": [float(c[1]) for c in sorted_coords],
        "z": [float(c[2]) for c in sorted_coords],
    })
    content = "\n".join(header) + "\n" + df.to_string(index=False, header=False)

    geom_path = out_base.with_suffix('.geom')
    geom_path.write_text(content, encoding='utf-8')
    return str(geom_path)

# ──────────────────────────────────────────────
# Open Babel helpers 

def have_obabel() -> bool:
    try:
        r = subprocess.run(["obabel", "-V"], stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        return r.returncode == 0
    except FileNotFoundError:
        print("Open Babel not installed")
        return False

def convert_with_obabel(log_path: str, mol2_path: str) -> None:
    r = subprocess.run(["obabel", log_path, "-O", mol2_path], stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
    if r.returncode != 0 or not os.path.isfile(mol2_path) or os.path.getsize(mol2_path) == 0:
        raise RuntimeError(f"Open Babel failed. STDOUT:{r.stdout} STDERR:{r.stderr}")

def load_mol_from_mol2(mol2_path: str) -> Chem.Mol:
    mol = Chem.MolFromMol2File(mol2_path, removeHs=False, sanitize=True)
    if mol is None:
        raise ValueError("No molecule could be read from .mol2.")
    return mol

# ──────────────────────────────────────────────
# Remove C-bound H (cleaner 2D/CDXML)

def remove_carbon_hydrogens(mol: Chem.Mol) -> Chem.Mol:
    rw = Chem.RWMol(mol)
    to_remove = []
    for atom in rw.GetAtoms():
        if atom.GetAtomicNum() == 1:
            nbrs = atom.GetNeighbors()
            if len(nbrs) == 1 and nbrs[0].GetAtomicNum() == 6:
                to_remove.append(atom.GetIdx())
    for idx in reversed(to_remove):
        rw.RemoveAtom(idx)
    mol2 = rw.GetMol()
    sanitize_ops = (
        Chem.SanitizeFlags.SANITIZE_ADJUSTHS |
        Chem.SanitizeFlags.SANITIZE_KEKULIZE |
        Chem.SanitizeFlags.SANITIZE_SETAROMATICITY
    )
    Chem.SanitizeMol(mol2, sanitizeOps=sanitize_ops)
    return mol2

# ──────────────────────────────────────────────
# 3D rendering (xyzrender)
DEFAULT_XYZRENDER_ARGS = "-S 1200"

def run_xyzrender(input_path: str, out_png: str, extra_args=None) -> None:
    """Run xyzrender with additional arguments (list or string)."""
    if not extra_args:
        extra_list = []
    elif isinstance(extra_args, str):
        extra_list = shlex.split(extra_args)
    else:
        try:
            extra_list = list(extra_args)
        except TypeError:
            extra_list = [str(extra_args)]
    cmd = ["xyzrender", input_path, "--output", out_png] + extra_list
    subprocess.run(cmd, check=True)

# ──────────────────────────────────────────────
# File cleanup helpers

def remove_file(path: str | None, verbose: bool = False, label: str = "file") -> None:
    if not path or not os.path.isfile(path):
        return
    try:
        os.remove(path)
        if verbose:
            print(f"[info] Removed intermediate {label}: {path}")
    except Exception as exc:
        print(f"[warning] Could not remove {path}: {exc}", file=sys.stderr)

def cleanup_glob(pattern: str = "*.v000.xyz", verbose: bool = False) -> int:
    removed = 0
    for path in glob.glob(pattern):
        try:
            os.remove(path)
            removed += 1
        except Exception as exc:
            print(f"[warning] Could not remove {path}: {exc}", file=sys.stderr)
    if verbose:
        print(f"[info] Removed {removed} file(s) matching '{pattern}'")
    return removed

# ──────────────────────────────────────────────
# CDXML generation + slide properties

def strip_stereo_rdkit(mol: Chem.Mol) -> Chem.Mol:
    # Make an editable copy
    rw = Chem.RWMol(mol)

    # Clear bond stereo and directions 
    for b in rw.GetBonds():
        b.SetStereo(BondStereo.STEREONONE)
        b.SetBondDir(BondDir.NONE)

    # Clear atom chirality
    for a in rw.GetAtoms():
        a.SetChiralTag(ChiralType.CHI_UNSPECIFIED)

    # Also clear any computed CIP labels if present
    for a in rw.GetAtoms():
        if a.HasProp('_CIPCode'):
            a.ClearProp('_CIPCode')

    new_mol = rw.GetMol()

    # Remove stereochemistry perception at the RDKit level
    Chem.RemoveStereochemistry(new_mol)
    return new_mol

def generate_cdxml(
    log_path: str,
    out_cdxml: str,
    all_cdxmls: list,
    all_props: list,
    ) -> tuple[str | None, bool, str | None]:
    used_obabel = False
    mol = None
    mol2_path = str(Path(out_cdxml).with_suffix('.mol2'))

    if have_obabel():
        try:
            convert_with_obabel(log_path, mol2_path)
            mol = load_mol_from_mol2(mol2_path)
            mol = strip_stereo_rdkit(mol)
            used_obabel = True
        except Exception as exc:
            print(f"[warning] HOpen Babel failed ({exc}).", file=sys.stderr)

    if mol is not None:
        # Clean C-H for nicer 2D
        try:
            mol = remove_carbon_hydrogens(mol)
        except Exception:
            pass

    if not HAVE_PYCDXML:
        print("[warning] pycdxml is not available; skipping CDXML and slide.", file=sys.stderr)
        return None, used_obabel, mol2_path if used_obabel else None

    if mol is None:
        # No Open Babel or no readable molecule -> cannot generate CDXML
        return None, used_obabel, mol2_path if used_obabel else None

    cdxml_doc = cdxml_converter.mol_to_document(mol)
    cdxml_converter.write_cdxml_file(cdxml_doc, out_cdxml)

    props_list: list = []
    base = os.path.splitext(os.path.basename(log_path))[0]
    props_list.append(cdxml_slide_generator.TextProperty('id', base, color='#3f6eba'))

    all_props.append(props_list)
    all_cdxmls.append(cdxml_doc.to_cdxml())

    return out_cdxml, used_obabel, mol2_path if used_obabel else None

# ──────────────────────────────────────────────
# DOCX assembly

def append_entry_to_docx(
    doc: Document,
    base_name: str,
    png_path: str,
    geom_text: str,
    title_size_pt: int,
    image_width_in: float,
    add_page_break: bool,
) -> None:
    p = doc.add_paragraph()
    run = p.add_run(base_name)
    run.bold = True
    run.font.size = Pt(title_size_pt)

    if png_path and os.path.isfile(png_path):
        pic_par = doc.add_paragraph()
        pic_par.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pic_par.add_run().add_picture(png_path, width=Inches(image_width_in))

    doc.add_paragraph(geom_text)

    if add_page_break:
        doc.add_page_break()

# ──────────────────────────────────────────────
# CLI parser

def build_parser() -> argparse.ArgumentParser:
    description = (
        "Gau2Publish: .log -> (.geom + .png + .cdxml) -> DOCX + optional ALL.cdxml"
    )
    epilog = (
        "Example:"
        " gau2publish *.log \
"        " --docx ESI_xyz.docx \
"        " --method wB97XD \
"        " --image-width 6 --title-size 14 \
"        " --slide-columns 6 \
"        " --xyzrender-args '-I -S 1200' \
"        " --rebuild \
"        " --scf-index 4,2"
    )
    parser = argparse.ArgumentParser(description=description, formatter_class=argparse.RawDescriptionHelpFormatter, epilog=epilog)

    parser.add_argument('logs', nargs='+', help='Gaussian .log files to process')
    parser.add_argument('-v', '--verbose', action='store_true', help='Keep intermediates and show extra messages')
    parser.add_argument('--version', action='version', version=f'%(prog)s {__version__}')

    grp_out = parser.add_argument_group('Main output')
    grp_out.add_argument('--docx', default='ESI_xyz.docx', help='Output DOCX filename (default: %(default)s)')
    grp_out.add_argument('--method', default='wB97XD', help='Method label for energy display (default: %(default)s)')
    grp_out.add_argument('--image-width', type=float, default=6.0, help='PNG width in inches (default: %(default)s)')
    grp_out.add_argument('--title-size', type=int, default=14, help='Title font size in pt (default: %(default)s)')

    grp_slide = parser.add_argument_group('ALL.cdxml mosaic')
    grp_slide.add_argument('--no-slide', action='store_true', help='Do not generate all.cdxml')
    grp_slide.add_argument('--slide-columns', type=int, default=6, help='Number of columns in the mosaic (default: %(default)s)')

    grp_xyz = parser.add_argument_group('3D rendering (xyzrender)')
    grp_xyz.add_argument('--xyzrender-args', default=DEFAULT_XYZRENDER_ARGS, help='Arguments forwarded to xyzrender (default: "%(default)s")')
    grp_xyz.add_argument('--no-xyzrender', action='store_true', help='Do not run xyzrender (no PNG will be generated)')

    rebuild_grp = parser.add_argument_group('Rebuild DOCX from existing files')
    rebuild_grp.add_argument('--rebuild', action='store_true', help='Only assemble DOCX from existing .geom and .png with the same basenames; do not render or generate CDXML; never delete PNGs')

    grp_idx = parser.add_argument_group('SCF energy selection')
    grp_idx.add_argument('--scf-index', metavar='base,high', help='Define indices: D_base=base+1, D_high=(base+1)-high; Python indices = -(D_*). Default: -1,-1')

    return parser

# ──────────────────────────────────────────────
# MAIN

def main() -> None:
    args = build_parser().parse_args()

    doc = Document()
    all_cdxmls: list = []
    all_props: list = []

    # --rebuild: assemble DOCX only from existing .geom/.png
    if args.rebuild:
        added = 0
        for item in args.logs:
            base = os.path.splitext(os.path.basename(item))[0]
            geom_path = str((Path(os.getcwd()) / base).with_suffix('.geom'))
            png_path = str((Path(os.getcwd()) / base).with_suffix('.png'))

            if not os.path.isfile(geom_path):
                print(f"[warning] Missing .geom for base '{base}': {geom_path}", file=sys.stderr)
                continue
            try:
                geom_text = Path(geom_path).read_text(encoding='utf-8', errors='replace')
            except Exception as exc:
                print(f"[warning] Could not read {geom_path}: {exc}", file=sys.stderr)
                continue

            append_entry_to_docx(
                doc, base,
                png_path if os.path.isfile(png_path) else '',
                geom_text,
                args.title_size,
                args.image_width,
                add_page_break=False
            )
            added += 1
            print(f"[ok] Added from existing files: {base}")

        out_docx = os.path.abspath(args.docx)
        doc.save(out_docx)
        print(f"[ok] DOCX saved: {out_docx}")
        return

    for idx, log_path in enumerate(args.logs):
        log_path = os.path.abspath(log_path)
        if not os.path.isfile(log_path):
            print(f"[error] Not found: {log_path}", file=sys.stderr)
            continue

        base = os.path.splitext(os.path.basename(log_path))[0]
        out_base = Path(os.getcwd()) / base

        # 1) .log -> data -> .geom
        try:
            (
                charge, mult, scf_energy, scf_energy_high, enthalpy,
                free_energy, free_energy_corr, frequencies,
                atomic_numbers, geometry
            ) = read_gaussian_output(log_path, scf_index=args.scf_index)
        except Exception as exc:
            print(f"[error] Could not parse {log_path}: {exc}", file=sys.stderr)
            continue

        geom_file = write_geom_file(
            charge, mult, scf_energy, scf_energy_high, free_energy_corr,
            frequencies, atomic_numbers, geometry, args.method, out_base
        )

        # 2) 3D render using xyzrender (PNG)
        png_path = str(out_base.with_suffix('.png'))
        png_created = False
        if not args.no_xyzrender:
            try:
                extra = args.xyzrender_args if isinstance(args.xyzrender_args, str) else args.xyzrender_args
                if args.verbose:
                    print(f"[info] xyzrender args: {shlex.split(extra) if isinstance(extra, str) else extra}")
                run_xyzrender(log_path, png_path, extra)
                png_created = os.path.isfile(png_path)
                cleanup_glob("*.v000.xyz", verbose=args.verbose)
            except Exception as exc:
                print(f"[warning] xyzrender unavailable or failed: {exc}", file=sys.stderr)
        else:
            png_path = ''

        # 3) CDXML generation
        cdxml_path = str(out_base.with_suffix('.cdxml'))
        mol2_path = None
        try:
            _, _, mol2_path = generate_cdxml(
                log_path, cdxml_path, all_cdxmls, all_props)
        except Exception as exc:
            print(f"[warning] CDXML not generated for {base}: {exc}", file=sys.stderr)

        # 4) Append to DOCX
        try:
            geom_text = Path(geom_file).read_text(encoding='utf-8', errors='replace')
        except Exception:
            geom_text = ''

        append_entry_to_docx(
            doc, base, png_path, geom_text,
            args.title_size, args.image_width,
            add_page_break=(idx < len(args.logs) - 1)
        )

        # 5) Cleanup intermediates if not verbose (includes .cdxml)
        if not args.verbose:
            remove_file(geom_file, verbose=False, label='.geom')
            remove_file(png_path if png_created else None, verbose=False, label='.png')
            remove_file(mol2_path, verbose=False, label='.mol2')
            remove_file(cdxml_path if os.path.isfile(cdxml_path) else None, verbose=False, label='.cdxml')
        else:
            print("[info] Intermediates kept.")

        print(f"[ok] Processed: {base}")

    # Save DOCX
    out_docx = os.path.abspath(args.docx)
    doc.save(out_docx)
    print(f"[ok] DOCX saved: {out_docx}")

    # Optional ALL.cdxml
    if not args.no_slide and HAVE_PYCDXML and all_cdxmls:
        try:
            n_props = len(all_props[0]) if all_props and all_props[0] else 1
            sg = cdxml_slide_generator.CDXMLSlideGenerator(style='ACS 1996', number_of_properties=n_props, columns=args.slide_columns)
            slide = sg.generate_slide(all_cdxmls, all_props)
            with open("all.cdxml", 'w', encoding='UTF-8') as fh:
                fh.write(slide)
            print('[ok] all.cdxml generated.')
        except Exception as exc:
            print(f"[warning] Slide generation failed: {exc}", file=sys.stderr)
    elif not args.no_slide and not HAVE_PYCDXML:
        print('[warning] Slide requested but pycdxml is not available.', file=sys.stderr)
